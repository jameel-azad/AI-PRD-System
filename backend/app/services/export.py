"""
PRD export service — converts a PRDVersion content dict to Markdown, PDF, or DOCX.

Each function accepts:
  prd_content : dict   — the PRDVersion.content JSON (section_key -> {content, ...})
  project_name: str    — shown in the title / cover page

export_prd() is the unified entry point; it returns (bytes, content_type, filename).
"""
from __future__ import annotations

import io
import logging
import re
from datetime import date

logger = logging.getLogger(__name__)

_SECTION_LABELS: dict[str, str] = {
    "project_overview":            "Project Overview",
    "business_objectives":         "Business Objectives",
    "stakeholders_personas":       "Stakeholders & Personas",
    "scope":                       "Scope",
    "functional_requirements":     "Functional Requirements",
    "non_functional_requirements": "Non-Functional Requirements",
    "user_stories":                "User Stories",
    "technical_constraints":       "Technical Constraints",
    "data_requirements":           "Data Requirements",
    "timeline_milestones":         "Timeline & Milestones",
    "assumptions_dependencies":    "Assumptions & Dependencies",
    "glossary":                    "Glossary",
    "open_questions":              "Open Questions",
    "source_index":                "Source Index",
}

_SECTION_ORDER = list(_SECTION_LABELS.keys())


def _label(key: str) -> str:
    return _SECTION_LABELS.get(key, key.replace("_", " ").title())


def _section_text(data) -> str:
    if isinstance(data, dict):
        return data.get("content", "")
    if isinstance(data, list):
        return "\n".join(
            item.get("question", str(item)) if isinstance(item, dict) else str(item)
            for item in data
        )
    return str(data) if data else ""


def _iter_sections(prd_content: dict):
    """Yield (key, label, text) in canonical display order."""
    seen: set[str] = set()
    for key in _SECTION_ORDER:
        if key in prd_content:
            yield key, _label(key), _section_text(prd_content[key])
            seen.add(key)
    for key, val in prd_content.items():
        if key not in seen and not key.startswith("_"):
            yield key, _label(key), _section_text(val)


# ── Shared text utilities ──────────────────────────────────────────────────────

def _split_source(text: str) -> tuple[str, str]:
    """Return (clean_text, sources_string) by extracting [Source: ...] citations.

    Deduplicates identical citations so the same file/timestamp doesn't repeat.
    """
    sources: list[str] = []
    seen_sources: set[str] = set()

    def _pull(m: re.Match) -> str:
        raw = m.group(1).strip()
        # Each citation may itself be pipe-separated (multi-source)
        for part in re.split(r"\s*\|\s*", raw):
            part = part.strip()
            if part and part not in seen_sources:
                seen_sources.add(part)
                sources.append(part)
        return ""

    clean = re.sub(r"\[Source:\s*([^\]]+)\]", _pull, text).strip()
    # Trim trailing punctuation left by citation removal
    clean = clean.rstrip(" .").strip()
    return clean, " | ".join(sources)


def _normalize_section_text(text: str) -> str:
    """Split inline-citation prose into individual bullet lines.

    The LLM sometimes emits content as a continuous paragraph where each
    sentence is followed by its [Source:] citation.  This function converts:
        'Req A. [Source: x] Req B. [Source: y]'
    into:
        '- Req A. [Source: x]\n- Req B. [Source: y]'

    Lines that already start with a list marker are left untouched.
    """
    result: list[str] = []
    for raw in text.split("\n"):
        line = raw.strip()
        if not line:
            result.append("")
            continue
        # Already a structured line — keep as-is
        if re.match(r"^[-*•#]\s+|^\d+[.)]\s+", line):
            result.append(line)
            continue
        # Count inline citations
        cite_count = len(re.findall(r"\[Source:", line))
        if cite_count <= 1:
            result.append(line)
            continue
        # Multiple citations: split at citation boundaries into individual items
        # Split on the closing ] followed by optional period/space and an uppercase letter
        parts = re.split(r"(\[Source:[^\]]+\])", line)
        buf = ""
        for part in parts:
            if re.match(r"\[Source:", part):
                buf = buf.strip().rstrip(".")
                if buf:
                    result.append(f"- {buf} {part}".strip())
                buf = ""
            else:
                buf += part
        # Trailing text without a citation
        buf = buf.strip().rstrip(".")
        if buf:
            result.append(f"- {buf}")
    return "\n".join(result)


def _safe(text: str) -> str:
    """Transliterate common Unicode punctuation to ASCII for fpdf2 core fonts."""
    table = {
        "—": "-",    # em dash
        "–": "-",    # en dash
        "‘": "'",    # left single quote
        "’": "'",    # right single quote
        "“": '"',    # left double quote
        "”": '"',    # right double quote
        "…": "...",  # ellipsis
        "•": "-",    # bullet
        "·": "-",    # middle dot
        " ": " ",    # non-breaking space
        "→": "->",   # right arrow
        "←": "<-",   # left arrow
        "≤": "<=",   # less-than-or-equal
        "≥": ">=",   # greater-than-or-equal
        "°": "deg",  # degree sign
        "±": "+/-",  # plus-minus
        "×": "x",    # multiplication sign
        "÷": "/",    # division sign
        "₹": "Rs.",  # Indian Rupee sign
        "é": "e",    # e acute (accented)
        "à": "a",    # a grave
        "è": "e",    # e grave
    }
    for ch, repl in table.items():
        text = text.replace(ch, repl)
    return text.encode("latin-1", errors="replace").decode("latin-1")


# ── Markdown ──────────────────────────────────────────────────────────────────

def export_markdown(prd_content: dict, project_name: str) -> bytes:
    lines: list[str] = [
        f"# {project_name}",
        "",
        "**Product Requirements Document**",
        f"*Generated: {date.today().isoformat()}*",
        "",
        "---",
        "",
    ]
    for _key, label, text in _iter_sections(prd_content):
        lines.append(f"## {label}")
        lines.append("")
        lines.append(text.strip() if text.strip() else "*No content generated for this section.*")
        lines.append("")
        lines.append("---")
        lines.append("")

    gaps = prd_content.get("_gaps", [])
    if gaps:
        lines.append("## Open Gap Questions")
        lines.append("")
        for i, g in enumerate(gaps, 1):
            section = g.get("section", "")
            question = g.get("question", str(g))
            priority = g.get("priority", "medium")
            lines.append(f"{i}. **[{priority.upper()}]** *{_label(section)}* - {question}")
        lines.append("")

    return "\n".join(lines).encode("utf-8")


# ── PDF ───────────────────────────────────────────────────────────────────────

def _pdf_source_line(pdf, source: str) -> None:
    """Render a source citation in small indented purple-grey text."""
    pdf.set_x(pdf.l_margin + 6)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(130, 110, 190)
    pdf.multi_cell(0, 4, _safe(f"Source: {source}"))
    pdf.set_text_color(50, 50, 60)
    pdf.ln(1)


def _render_section_body(pdf, text: str) -> None:
    """Render section content with proper bullet/numbered list formatting."""
    if not text or not text.strip():
        pdf.set_font("Helvetica", "I", 10)
        pdf.set_text_color(150, 150, 150)
        pdf.multi_cell(0, 6, "No content generated for this section.")
        return

    text = _normalize_section_text(text)

    for raw in text.split("\n"):
        line = raw.strip()
        if not line:
            pdf.ln(1.5)
            continue

        # Subheadings (## or ### prefix)
        hm = re.match(r"^#{1,3}\s+(.*)", line)
        if hm:
            clean, src = _split_source(hm.group(1))
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(40, 40, 50)
            pdf.multi_cell(0, 6.5, _safe(clean))
            if src:
                _pdf_source_line(pdf, src)
            pdf.ln(2)
            continue

        # Numbered list items (1. or 1) prefix)
        nm = re.match(r"^(\d+)[.)]\s+(.*)", line)
        if nm:
            clean, src = _split_source(nm.group(2))
            if not clean:
                continue
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(100, 80, 200)
            # Number badge
            num_text = f"{nm.group(1)}."
            pdf.set_x(pdf.l_margin)
            pdf.cell(8, 5.5, _safe(num_text))
            pdf.set_font("Helvetica", "", 10.5)
            pdf.set_text_color(50, 50, 60)
            pdf.multi_cell(0, 5.5, _safe(clean))
            if src:
                _pdf_source_line(pdf, src)
            pdf.ln(1)
            continue

        # Bullet list items (-, *, • prefix)
        bm = re.match(r"^[-*•]\s+(.*)", line)
        if bm:
            clean, src = _split_source(bm.group(1))
            if not clean:
                continue
            pdf.set_x(pdf.l_margin + 3)
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(124, 110, 230)
            pdf.cell(6, 5.5, "-")
            pdf.set_font("Helvetica", "", 10.5)
            pdf.set_text_color(50, 50, 60)
            pdf.multi_cell(0, 5.5, _safe(clean))
            if src:
                _pdf_source_line(pdf, src)
            pdf.ln(1)
            continue

        # Plain paragraph line
        clean, src = _split_source(line)
        if clean:
            pdf.set_font("Helvetica", "", 10.5)
            pdf.set_text_color(50, 50, 60)
            pdf.multi_cell(0, 5.5, _safe(clean))
            if src:
                _pdf_source_line(pdf, src)
            pdf.ln(1)


def export_pdf(prd_content: dict, project_name: str) -> bytes:
    try:
        from fpdf import FPDF
    except ImportError:
        logger.error("fpdf2 is not installed")
        raise RuntimeError("fpdf2 package is required for PDF export. Run: pip install fpdf2")

    class PRDPdf(FPDF):
        def header(self):
            if self.page_no() == 1:
                return
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(150, 150, 150)
            # Use ASCII dash to avoid latin-1 encoding issues with em dash
            self.cell(0, 8, _safe(f"Xccelera PRD Portal - {project_name}"), align="R")
            self.ln(2)
            self.set_draw_color(220, 220, 220)
            self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
            self.ln(4)

        def footer(self):
            if self.page_no() == 1:
                return
            self.set_y(-14)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 8, f"Page {self.page_no()}", align="C")

    pdf = PRDPdf(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=22)
    pdf.set_margins(left=20, top=20, right=20)

    # ── Cover page ────────────────────────────────────────────────────────────
    pdf.add_page()
    pdf.set_fill_color(15, 15, 20)
    pdf.rect(0, 0, 210, 297, "F")

    # Purple accent bar
    pdf.set_fill_color(124, 110, 230)
    pdf.rect(0, 108, 210, 4, "F")

    # Company label above title
    pdf.set_y(55)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(130, 110, 190)
    pdf.cell(0, 6, "XCCELERA PRD PORTAL", align="C")

    pdf.ln(8)
    pdf.set_font("Helvetica", "B", 30)
    pdf.set_text_color(255, 255, 255)
    pdf.multi_cell(0, 13, _safe(project_name), align="C")

    pdf.ln(6)
    pdf.set_font("Helvetica", "", 13)
    pdf.set_text_color(160, 148, 255)
    pdf.cell(0, 8, "Product Requirements Document", align="C")

    pdf.ln(22)
    pdf.set_font("Helvetica", "", 9.5)
    pdf.set_text_color(120, 120, 140)
    pdf.cell(0, 6, _safe(f"Generated: {date.today().strftime('%B %d, %Y')}"), align="C")

    # ── Table of Contents ─────────────────────────────────────────────────────
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(15, 15, 20)
    pdf.cell(0, 10, "Table of Contents", ln=True)
    pdf.set_draw_color(124, 110, 230)
    pdf.set_line_width(0.5)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(8)

    # Collect sections with page estimates (start at page 3)
    page_num = 3
    toc_entries: list[tuple[str, int]] = []
    for _key, label, text in _iter_sections(prd_content):
        toc_entries.append((label, page_num))
        # Rough page estimate: 1 base + 1 per ~1500 chars
        page_num += max(1, len(text) // 1500)

    gaps = prd_content.get("_gaps", [])
    if gaps:
        toc_entries.append(("Open Gap Questions", page_num))

    for i, (entry_label, entry_page) in enumerate(toc_entries, 1):
        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(50, 50, 60)
        num_str = f"{i:02d}."
        pdf.cell(10, 7, _safe(num_str))
        pdf.cell(0, 7, _safe(entry_label), ln=True)

    # ── Sections ──────────────────────────────────────────────────────────────
    for _key, label, text in _iter_sections(prd_content):
        pdf.add_page()

        # Section heading
        pdf.set_font("Helvetica", "B", 17)
        pdf.set_text_color(15, 15, 20)
        pdf.cell(0, 10, _safe(label), ln=True)

        pdf.set_draw_color(124, 110, 230)
        pdf.set_line_width(0.6)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(7)

        _render_section_body(pdf, text)

    # ── Gap questions ─────────────────────────────────────────────────────────
    if gaps:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 17)
        pdf.set_text_color(15, 15, 20)
        pdf.cell(0, 10, "Open Gap Questions", ln=True)
        pdf.set_draw_color(230, 180, 50)
        pdf.set_line_width(0.6)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(7)

        for i, g in enumerate(gaps, 1):
            section = _label(g.get("section", ""))
            question = g.get("question", str(g))
            priority = g.get("priority", "medium").upper()

            # Priority badge
            badge_color = (180, 40, 40) if priority == "HIGH" else (200, 140, 20) if priority == "MEDIUM" else (80, 140, 80)
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(*badge_color)
            pdf.cell(0, 5.5, _safe(f"{i}.  [{priority}]  {section}"), ln=True)

            pdf.set_font("Helvetica", "", 10.5)
            pdf.set_text_color(50, 50, 60)
            pdf.set_x(pdf.l_margin + 6)
            pdf.multi_cell(0, 5.5, _safe(question))
            pdf.ln(4)

    return bytes(pdf.output())


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _render_section_docx(doc, text: str) -> None:
    """Add paragraphs to a DOCX document with proper list formatting."""
    from docx.shared import Pt, RGBColor

    if not text or not text.strip():
        p = doc.add_paragraph()
        run = p.add_run("No content generated for this section.")
        run.italic = True
        run.font.color.rgb = RGBColor(0x96, 0x96, 0x96)
        return

    text = _normalize_section_text(text)

    for raw in text.split("\n"):
        line = raw.strip()
        if not line:
            continue

        # Subheadings
        hm = re.match(r"^#{1,3}\s+(.*)", line)
        if hm:
            clean, src = _split_source(hm.group(1))
            h = doc.add_heading(clean, level=2)
            h.runs[0].font.size = Pt(12)
            h.runs[0].font.color.rgb = RGBColor(0x2D, 0x2D, 0x3E)
            if src:
                sp = doc.add_paragraph()
                sr = sp.add_run(f"Source: {src}")
                sr.italic = True
                sr.font.size = Pt(8)
                sr.font.color.rgb = RGBColor(0x82, 0x6E, 0xBE)
            continue

        # Numbered list
        nm = re.match(r"^(\d+)[.)]\s+(.*)", line)
        if nm:
            clean, src = _split_source(nm.group(2))
            if not clean:
                continue
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(clean)
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x32, 0x32, 0x3C)
            if src:
                src_run = p.add_run(f"  [Source: {src}]")
                src_run.italic = True
                src_run.font.size = Pt(8.5)
                src_run.font.color.rgb = RGBColor(0x82, 0x6E, 0xBE)
            continue

        # Bullet list
        bm = re.match(r"^[-*•]\s+(.*)", line)
        if bm:
            clean, src = _split_source(bm.group(1))
            if not clean:
                continue
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(clean)
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x32, 0x32, 0x3C)
            if src:
                src_run = p.add_run(f"  [Source: {src}]")
                src_run.italic = True
                src_run.font.size = Pt(8.5)
                src_run.font.color.rgb = RGBColor(0x82, 0x6E, 0xBE)
            continue

        # Plain paragraph
        clean, src = _split_source(line)
        if clean:
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x32, 0x32, 0x3C)
            if src:
                src_run = p.add_run(f"  [Source: {src}]")
                src_run.italic = True
                src_run.font.size = Pt(8.5)
                src_run.font.color.rgb = RGBColor(0x82, 0x6E, 0xBE)


def export_docx(prd_content: dict, project_name: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx is not installed")
        raise RuntimeError("python-docx package is required for DOCX export. Run: pip install python-docx")

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # Cover
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(project_name)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x0F, 0x0F, 0x14)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Product Requirements Document")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x7C, 0x6E, 0xE6)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {date.today().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x80, 0x80, 0x90)

    doc.add_page_break()

    # Sections
    for _key, label, text in _iter_sections(prd_content):
        heading = doc.add_heading(label, level=1)
        heading.runs[0].font.color.rgb = RGBColor(0x0F, 0x0F, 0x14)
        _render_section_docx(doc, text)
        doc.add_paragraph()  # spacer

    # Gap questions
    gaps = prd_content.get("_gaps", [])
    if gaps:
        doc.add_page_break()
        heading = doc.add_heading("Open Gap Questions", level=1)
        heading.runs[0].font.color.rgb = RGBColor(0xD9, 0x7B, 0x06)

        for i, g in enumerate(gaps, 1):
            section = _label(g.get("section", ""))
            question = g.get("question", str(g))
            priority = g.get("priority", "medium").upper()

            q_para = doc.add_paragraph()
            label_run = q_para.add_run(f"{i}. [{priority}] {section} - ")
            label_run.bold = True
            label_run.font.color.rgb = RGBColor(0x50, 0x3C, 0xB4)
            q_para.add_run(question)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Unified entry point ───────────────────────────────────────────────────────

def export_prd(
    prd_content: dict,
    project_name: str,
    fmt: str,
) -> tuple[bytes, str, str]:
    """
    Export a PRD in the requested format.
    Returns (data_bytes, content_type, filename).
    """
    fmt = fmt.lower().strip()
    safe_name = re.sub(r"[^a-zA-Z0-9_-]", "_", project_name)[:60]

    if fmt in ("md", "markdown"):
        data = export_markdown(prd_content, project_name)
        return data, "text/markdown; charset=utf-8", f"{safe_name}_PRD.md"

    if fmt == "pdf":
        data = export_pdf(prd_content, project_name)
        return data, "application/pdf", f"{safe_name}_PRD.pdf"

    if fmt == "docx":
        data = export_docx(prd_content, project_name)
        return data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{safe_name}_PRD.docx"

    raise ValueError(f"Unsupported export format: {fmt!r}. Use 'pdf', 'docx', or 'md'.")
